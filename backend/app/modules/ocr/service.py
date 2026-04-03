"""
OCR Service (Diagram-Aware)

Pipeline:
  1. Preprocess PDF pages into images.
  2. Unified Qwen-Vision call per page — extracts text AND produces a structured
     evaluation description for each diagram inline using <diagram> … </diagram> tags.
  3. Build a rich .docx:
       • Page headings
       • Plain text rendered as normal paragraphs
       • Each <diagram> block rendered as a styled, labelled evaluation section
  4. PostProcessor slices full text using a two-phase sliding window → JSON
     with fields: question_id, answer, diagram_present, diagram_description.
"""

from __future__ import annotations

import os
import re
from typing import List, Tuple, Dict, Any

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.modules.ocr.preprocessor import ImagePreprocessor
from app.modules.ocr.extractor import TextExtractor
from app.modules.ocr.postprocessor import PostProcessor


# ── Regex that splits the structured page text into alternating text / diagram chunks ──
_SPLIT_RE = re.compile(
    r'(<diagram[^>]*>.*?</diagram>)',
    re.DOTALL | re.IGNORECASE,
)

# ── Regex to pull the optional index= attribute so we can label the box ────
_INDEX_RE = re.compile(r'index=["\']?(\d+)["\']?', re.IGNORECASE)


class OCRService:
    """
    Orchestrates the OCR pipeline for handwritten answer sheet processing.

    Usage:
        service = OCRService()
        result = await service.process("path/to/file.pdf", "application/pdf")
        # result = {"Q1": "answer text", "Q2": "answer text", ...}
    """

    def __init__(self):
        self.preprocessor  = ImagePreprocessor()
        self.extractor     = TextExtractor()
        self.postprocessor = PostProcessor()

    # ------------------------------------------------------------------ #
    #  Main entry point                                                    #
    # ------------------------------------------------------------------ #

    async def process(self, file_path: str, file_type: str) -> List[Dict[str, Any]]:
        pdf_name   = os.path.splitext(os.path.basename(file_path))[0]
        output_dir = os.path.join(os.path.dirname(file_path), "OCR_Result")
        os.makedirs(output_dir, exist_ok=True)
        output_doc_path = os.path.join(output_dir, pdf_name + ".docx")

        image_paths: List[str] = []
        try:
            # Step 1 — Preprocess
            print(f"\nPreprocessing {pdf_name}...")
            image_paths, _ = await self.preprocessor.preprocess(file_path, file_type)

            # Step 2 — Unified OCR + diagram extraction
            print(f"Running unified OCR + diagram extraction on "
                  f"{len(image_paths)} page(s)...")
            page_results: List[Tuple[str, Dict]] = await self.extractor.extract(image_paths)

        finally:
            for path in image_paths:
                try:
                    os.remove(path)
                except OSError:
                    pass

        # Step 3 — Build rich DOCX
        print("Building DOCX...")
        self._build_docx(output_doc_path, pdf_name, page_results)
        print(f"\nOCR Completed. Saved: {output_doc_path}")

        # Step 4 — PostProcessor → JSON
        json_path = self.postprocessor.process(output_doc_path, page_results)
        print(f"Q&A JSON saved: {json_path}")

        import json
        with open(json_path, 'r', encoding='utf-8') as f:
            records = json.load(f)

        return records

    # ------------------------------------------------------------------ #
    #  DOCX builder                                                        #
    # ------------------------------------------------------------------ #

    def _build_docx(
        self,
        output_path: str,
        pdf_name: str,
        page_results: List[Tuple[str, Dict]],
    ) -> None:
        doc = Document()
        doc.add_heading(f"OCR Result for {pdf_name}", level=1)

        for i, (structured_text, meta) in enumerate(page_results):
            doc.add_heading(f"Page {i + 1}", level=2)

            # Split the structured text into plain-text chunks and diagram blocks
            chunks = _SPLIT_RE.split(structured_text)

            for chunk in chunks:
                chunk = chunk.strip()
                if not chunk:
                    continue

                if re.match(r'<diagram', chunk, re.IGNORECASE):
                    self._add_diagram_block(doc, chunk)
                else:
                    # Plain text — preserve line breaks
                    for line in chunk.splitlines():
                        line = line.strip()
                        if line:
                            doc.add_paragraph(line)

        doc.save(output_path)

    # ------------------------------------------------------------------ #
    #  Diagram block renderer                                              #
    # ------------------------------------------------------------------ #

    def _add_diagram_block(self, doc_obj: Any, raw_block: str) -> None:
        """
        Render a <diagram …> … </diagram> block as a visually distinct section:
          ┌──────────────────────────────────────┐
          │  [Diagram N — Evaluation]            │  ← bold, shaded heading paragraph
          │  <structured description text>       │  ← monospace paragraphs
          └──────────────────────────────────────┘
        """
        # Extract the inner content
        inner_match = re.search(
            r'<diagram[^>]*>(.*?)</diagram>', raw_block, re.DOTALL | re.IGNORECASE
        )
        inner = inner_match.group(1).strip() if inner_match else raw_block

        # Extract index for the heading label
        idx_match = _INDEX_RE.search(raw_block)
        label = f"Diagram {idx_match.group(1)} — Evaluation" if idx_match else "Diagram — Evaluation"

        # ── Heading paragraph ──
        heading_para = doc_obj.add_paragraph()
        heading_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = heading_para.add_run(f"[ {label} ]")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1A, 0x53, 0x76)   # dark blue

        # ── Content paragraphs (monospace) ──
        for line in inner.splitlines():
            line_para = doc_obj.add_paragraph(line)
            for run in line_para.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(9)

        # Blank separator
        doc_obj.add_paragraph()